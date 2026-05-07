"""PlaySmart Phase 4 Report DOCX Generator"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

BASE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(BASE, "hifi-screenshots")
OUT = os.path.join(BASE, "PlaySmart-Phase4-Report.docx")

def add_hyperlink(p, url, text):
    part = p.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl = parse_xml(f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>')
    r = parse_xml(f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r>')
    hl.append(r); p._element.append(hl)

def shade(cell, color):
    tc = cell._element.get_or_add_tcPr()
    tc.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'))

def table_with_header(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; shade(c,"2E75B6")
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c = t.rows[1+ri].cells[ci]; c.text = v
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
            if ri%2==1: shade(c,"D9E2F3")
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    return t

def add_img(doc, path, w=5.0, caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f"[Image: {os.path.basename(path)}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].italic=True; c.runs[0].font.size=Pt(9); c.runs[0].font.color.rgb=RGBColor(0x66,0x66,0x66)

def build():
    doc = Document()
    s = doc.sections[0]; s.top_margin=s.bottom_margin=s.left_margin=s.right_margin=Inches(1)
    style = doc.styles["Normal"]; style.font.name="Calibri"; style.font.size=Pt(11)

    # TITLE PAGE
    for _ in range(5): doc.add_paragraph("")
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=t.add_run("PlaySmart"); r.bold=True; r.font.size=Pt(36); r.font.color.rgb=RGBColor(0x1A,0x5C,0x97)
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=t.add_run("Phase 4: High-Fidelity Design and Heuristic Evaluation"); r.font.size=Pt(18); r.font.color.rgb=RGBColor(0x44,0x44,0x44)
    doc.add_paragraph("")
    for line in ["Team Members: Usman Wajid, Ahmad Masood, Abdulrehman Naseer","Instructor: Arslan Asif","Date: April 2025"]:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(line); r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x55,0x55,0x55)
    doc.add_page_break()

    # TOC
    doc.add_heading("Table of Contents",1)
    for item in ["1. Introduction","2. Design System","3. High-Fidelity Screen Catalog","4. Navigation Flow & Interaction Design","5. UI Components Showcase","6. Design Principles Applied","7. Conclusion","Appendix: Design Links"]:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 1. INTRODUCTION
    doc.add_heading("1. Introduction",1)
    doc.add_paragraph("This phase represents the culmination of the PlaySmart design process, transforming the low-fidelity wireframes from Phase 3 into polished, high-fidelity mockups. The designs have been created using Google Stitch and follow the \"Kinetic Precision\" design system - a dark, performance-focused visual language designed for athletes and coaches.")
    doc.add_paragraph("Phase 4 delivers two key artifacts:")
    doc.add_paragraph("High-Fidelity Designs: 14 complete screens covering all navigational paths, including player flows, coach flows, and system interaction states (dialogs, errors, progress indicators)", style="List Bullet")
    doc.add_paragraph("Heuristic Evaluation Workbook: A systematic evaluation of the designs against Jakob Nielsen's 10 Heuristic Principles, identifying strengths and areas for improvement", style="List Bullet")
    p=doc.add_paragraph(); r=p.add_run("Interactive Prototype: "); r.bold=True
    add_hyperlink(p,"https://stitch.withgoogle.com/projects/12363635124012461076","PlaySmart Hi-Fi Designs - Stitch")
    doc.add_page_break()

    # 2. DESIGN SYSTEM
    doc.add_heading("2. Design System",1)
    doc.add_heading("2.1 Color Palette",2)
    doc.add_paragraph("The \"Kinetic Precision\" design system uses a dark, high-contrast palette optimized for athletic environments:")
    table_with_header(doc,["Token","Hex","Usage"],[
        ("Surface (Base)","#0E0E0E","Primary background - charcoal foundation"),
        ("Primary (Volt)","#CAFD00","Critical actions, CTAs, active states"),
        ("Primary Light","#F3FFCA","Text on dark, subtle highlights"),
        ("Secondary","#ECE856","Supporting metrics, secondary indicators"),
        ("Tertiary","#FCE047","Warnings, tertiary accents"),
        ("Error","#FF7351","Error states, destructive actions"),
        ("Surface Container","#1A1A1A","Card backgrounds, elevated surfaces"),
        ("On Surface","#FFFFFF","Primary text on dark backgrounds"),
        ("On Surface Variant","#ADAAAA","Secondary text, metadata"),
        ("Outline","#767575","Borders, dividers"),
    ],[1.8,1.0,3.7])

    doc.add_heading("2.2 Typography",2)
    table_with_header(doc,["Role","Font","Usage"],[
        ("Display / Headline","Space Grotesk","Screen titles, large performance stats, drill names"),
        ("Title","Inter Semi-Bold","Card headers, section titles"),
        ("Body","Inter Regular","Descriptions, instructions, content text"),
        ("Label","Inter All-Caps","Metadata tags (REPS, SETS, BPM), timestamps"),
    ],[1.8,1.5,3.2])

    doc.add_heading("2.3 Layout Principles",2)
    for item in ["Tonal Layering: Depth through background color shifts, not drop shadows","No-Line Rule: Sections separated by tonal shifts, not 1px borders","Glassmorphism: Floating elements use backdrop-blur with semi-transparent backgrounds","Spacing Scale: Consistent 8px grid with generous whitespace for breathing room"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2.4 Shape & Roundness",2)
    doc.add_paragraph("Corner radius: 0.375rem (md) for buttons and cards. Full roundness for chips and action filters. Square corners for input fields and structural containers.")
    doc.add_page_break()

    # 3. SCREEN CATALOG
    doc.add_heading("3. High-Fidelity Screen Catalog",1)
    doc.add_paragraph("The following 14 screens represent the complete PlaySmart interface. Each screen has been designed with realistic content, consistent styling, and clear navigational paths.")

    screens = [
        ("welcome_and_login_screens","Welcome & Login Screens","The Welcome screen introduces PlaySmart with the tagline \"Train Smarter, Play Better\" and a prominent Get Started CTA. The Login screen provides email/password authentication with social login options and account creation."),
        ("player_dashboard_and_login_detail","Player Dashboard","The central hub for players showing quick actions (Plan Training, Drill Library, My Progress, Match Prep), upcoming session card, weekly progress bar, and motivational streak counter."),
        ("training_categories_and_weekly_progress","Training Categories","Organized drill categories (Shooting, Speed, Tactics) with visual cards. Each category displays drill count, estimated duration, and difficulty level."),
        ("speed_detail_and_training_progress","Speed Training Detail & In-Progress","Detailed drill view showing step-by-step exercises with rep counts. The in-progress screen displays a real-time timer, current exercise, and progress indicator."),
        ("training_completed_and_drill_library","Training Completed & Drill Library","Post-session summary showing total time, exercises completed, calories burned, and RPE effort rating (1-10 scale). The Drill Library provides a filterable grid of all available drills with category chips."),
        ("player_progress_tracking_and_ui_states","Player Progress Tracking","Visual dashboard with weekly/monthly toggle, circular progress ring (75% goal), bar chart showing daily activity, session statistics (12 sessions, 8.5 hours, 47 drills), and 5-day streak tracker."),
        ("coach_dashboard_and_calendar","Coach Dashboard & Training Calendar","Coach-specific home screen with Create Session, Player Roster, Drill Library, Session History quick actions. The Training Calendar shows weekly view with scheduled sessions color-coded by type."),
        ("match_preparation_and_calendar_bottom","Match Preparation","Pre-match checklist with completion toggles (Fitness Check, Tactical Drills, Warm-up Plan, Mental Readiness), team readiness overview with player fatigue indicators (green/yellow/red), and readiness scores."),
        ("profile_settings_and_final_ui_states","Profile & Settings / UI Interaction States","User profile management with notification toggles, theme preferences, and account options. The UI States screen showcases confirmation dialogs, success/error toast messages, and loading progress indicators."),
    ]

    for fname, title, desc in screens:
        doc.add_heading(title, 3)
        path = None
        for f in os.listdir(SHOTS):
            if fname in f and f.endswith('.png'):
                path = os.path.join(SHOTS, f); break
        if path: add_img(doc, path, 5.5, title)
        doc.add_paragraph(desc)
        doc.add_paragraph("")

    doc.add_heading("Project Overview",3)
    for prefix in ["playsmart_hifi_overview_top","playsmart_hifi_overview_bottom"]:
        for f in os.listdir(SHOTS):
            if prefix in f and f.endswith('.png'):
                add_img(doc, os.path.join(SHOTS, f), 6.0)
    doc.add_page_break()

    # 4. NAVIGATION FLOW
    doc.add_heading("4. Navigation Flow & Interaction Design",1)
    doc.add_heading("4.1 Bottom Navigation Bar",2)
    doc.add_paragraph("All screens share a consistent bottom navigation bar with 5 tabs:")
    table_with_header(doc,["Tab","Icon","Destination","Description"],[
        ("Home","House","Player/Coach Dashboard","Central hub with quick actions and overview"),
        ("Drills","Grid","Drill Library","Browse and filter all available drills"),
        ("Plan","Calendar","Training Calendar","Weekly schedule and session management"),
        ("Progress","Chart","Player Progress","Performance tracking and statistics"),
        ("Profile","Person","Settings & Profile","Account management and preferences"),
    ],[1.0,0.8,1.5,3.2])

    doc.add_heading("4.2 Key Navigation Paths",2)
    paths = [
        "Player Training Flow: Home > Training Categories > Speed/Shooting/Tactics > Drill Detail > Start Training > In-Progress Timer > Training Completed",
        "Coach Session Flow: Coach Dashboard > Create Session > Select Date > Choose Focus > Pick Drills > Review Plan > Session Created",
        "Drill Discovery: Home > Drill Library > Filter (All/Fitness/Skills/Tactical) > Search > Drill Card > Drill Detail",
        "Progress Review: Home > My Progress > Weekly/Monthly Toggle > View Charts > Session History",
        "Match Prep Flow: Home > Match Preparation > Review Checklist > Team Readiness > Start Prep Session",
    ]
    for p in paths: doc.add_paragraph(p, style="List Bullet")

    doc.add_heading("4.3 Dialog & System Feedback",2)
    table_with_header(doc,["Component","Type","Example","Trigger"],[
        ("Confirmation Dialog","Modal Overlay","Delete Training Session?","Destructive action (delete, remove)"),
        ("Success Toast","Banner Notification","Session saved successfully!","Successful operation completion"),
        ("Error Message","Banner Notification","Failed to sync data","Network/system error"),
        ("Progress Indicator","Circular + Bar","Syncing training data... 65%","Data loading/syncing"),
        ("Loading State","Skeleton Screen","Content loading placeholders","Initial page load"),
    ],[1.5,1.3,2.0,1.7])
    doc.add_page_break()

    # 5. UI COMPONENTS
    doc.add_heading("5. UI Components Showcase",1)
    doc.add_heading("5.1 Buttons",2)
    table_with_header(doc,["Type","Fill","Text Color","Usage"],[
        ("Primary CTA","#CAFD00 (Volt)","#3A4A00 (Dark)","Start Training, Save Session, Get Started"),
        ("Secondary","#262626 (Surface)","#F3FFCA (Light)","View Details, Back to Home"),
        ("Ghost/Outline","Transparent","#FFFFFF","Cancel, Dismiss, Secondary actions"),
        ("Destructive","#FF7351 (Error)","#FFFFFF","Delete, Remove"),
    ],[1.3,1.5,1.5,2.2])

    doc.add_heading("5.2 Cards",2)
    doc.add_paragraph("Cards use surface-container-high (#20201F) backgrounds with no border lines. Content separation is achieved through 1.5rem vertical spacing. Drill cards display: image/icon area, drill name (Title font), duration and difficulty metadata (Label font).")

    doc.add_heading("5.3 Input Fields",2)
    doc.add_paragraph("Text inputs use surface-container-highest (#262626) background. Focus state switches to surface-bright (#2C2C2C) with a primary-colored ghost border. Labels appear above inputs in Label font (all-caps).")

    doc.add_heading("5.4 Chips & Filters",2)
    doc.add_paragraph("Action chips use full roundness with secondary-container (#636100) background. Active state switches to primary (#CAFD00). Used for category filtering (All, Fitness, Skills, Tactical) in the Drill Library.")

    doc.add_heading("5.5 Progress Indicators",2)
    doc.add_paragraph("Circular progress rings display weekly goal completion (e.g., 75%). Bar charts show daily activity distribution. Linear progress bars with primary-dim glow effect indicate real-time training progress.")
    doc.add_page_break()

    # 6. DESIGN PRINCIPLES
    doc.add_heading("6. Design Principles Applied",1)

    doc.add_heading("6.1 Shneiderman's Eight Golden Rules",2)
    rules = [
        ("Strive for Consistency","Consistent bottom navigation, card styles, button designs, and typography across all 14 screens. The Kinetic Precision design system ensures uniform visual language."),
        ("Seek Universal Usability","Supports both player and coach roles with role-specific dashboards. High contrast (WCAG AA) ensures readability in various lighting conditions including outdoor fields."),
        ("Offer Informative Feedback","Toast notifications for save/delete actions, progress indicators during sync, checkmark animations on task completion, real-time timer updates during training."),
        ("Design Dialogs to Yield Closure","Training Completed screen provides session summary. Confirmation dialogs for destructive actions. Session creation ends with a review and confirmation step."),
        ("Prevent Errors","Form validation on login/registration, disabled buttons until required fields are complete, confirmation dialogs before deleting sessions, RPE scale with descriptive labels."),
        ("Permit Easy Reversal of Actions","Back navigation on every screen, undo capability for drill removal, cancel buttons on all dialogs, swipe-to-go-back gesture support."),
        ("Keep Users in Control","Manual RPE input, customizable notification preferences, player/coach mode switching, direct access to all features via bottom navigation."),
        ("Reduce Short-Term Memory Load","Dashboard shows all quick actions, drill cards display key info (name, duration, difficulty) upfront, progress ring shows goal status at a glance, calendar provides visual schedule overview."),
    ]
    for title, desc in rules:
        p=doc.add_paragraph(); r=p.add_run(f"{title}: "); r.bold=True; p.add_run(desc)

    doc.add_heading("6.2 Norman's Design Principles",2)
    norman = [
        ("Visibility","All interactive elements are clearly visible. Bottom nav always present. Primary CTAs use high-contrast volt green (#CAFD00) against dark backgrounds."),
        ("Feedback","Immediate visual feedback: button press states, toast notifications, progress animations, loading spinners, and completion checkmarks."),
        ("Constraints","Calendar restricts date selection to valid ranges. RPE scale limited to 1-10. Category filters constrain drill browsing to relevant results."),
        ("Mapping","Navigation icons map directly to their functions (house=home, calendar=plan). Green/yellow/red fatigue indicators map to intuitive risk levels."),
        ("Consistency","Uniform component library across all screens. Same interaction patterns for similar actions. Consistent use of Space Grotesk for headings, Inter for body."),
        ("Affordance","Buttons look pressable with filled backgrounds. Toggle switches clearly indicate on/off states. Cards appear tappable with subtle elevation cues."),
    ]
    for title, desc in norman:
        p=doc.add_paragraph(); r=p.add_run(f"{title}: "); r.bold=True; p.add_run(desc)

    doc.add_heading("6.3 Theo Mandel's Guidelines",2)
    mandel = [
        ("Place Users in Control","Users choose their own training path, can skip onboarding, switch between player/coach modes, and navigate freely via bottom tab bar."),
        ("Reduce Users' Memory Load","Dashboard quick-action cards eliminate need to remember navigation paths. Drill cards show all essential info. Calendar provides visual schedule reference."),
        ("Make the Interface Consistent","Single design system applied across all screens. Consistent spacing, typography, and color usage. Predictable interaction patterns throughout."),
    ]
    for title, desc in mandel:
        p=doc.add_paragraph(); r=p.add_run(f"{title}: "); r.bold=True; p.add_run(desc)

    doc.add_heading("6.4 Five Principles of Screen Design",2)
    table_with_header(doc,["Principle","Application in PlaySmart"],[
        ("Grouping","Related elements grouped in cards (e.g., session info, player stats). Quick actions in 2x2 grid. Settings organized by category (Notifications, Preferences, Account)."),
        ("Ordering","Screens follow logical task flow. Dashboard prioritizes most-used actions. Drill library sorts by category then difficulty. Calendar orders sessions chronologically."),
        ("Alignment","Strong left alignment for text content. Center alignment for progress rings and CTAs. Grid-based layout ensures consistent element positioning across screens."),
        ("Use of Whitespace","Generous padding (2rem+) between sections. No crowded layouts. Tonal layering creates visual breathing room without explicit separators."),
        ("Decoration","Minimal decorative elements. Volt green accent used purposefully for active states and CTAs only. No gratuitous icons or imagery. Content-first approach."),
    ],[1.5,5.0])
    doc.add_page_break()

    # 7. CONCLUSION
    doc.add_heading("7. Conclusion",1)
    doc.add_paragraph("Phase 4 demonstrates a clear and deliberate progression from the low-fidelity wireframes of Phase 3 to polished, production-ready high-fidelity designs. The PlaySmart interface now comprises 14 complete screens covering all user flows for both players and coaches, with comprehensive UI component coverage including dialogs, error states, and progress indicators.")
    doc.add_paragraph("Key achievements in this phase:")
    for item in ["14 high-fidelity screens with consistent Kinetic Precision design system","Complete navigation coverage with all paths documented","Role-based interfaces for both players and coaches","Comprehensive UI component library (buttons, cards, inputs, chips, dialogs, toasts)","Systematic application of Shneiderman's, Norman's, and Mandel's design principles","Thorough heuristic evaluation identifying strengths and improvement areas"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()

    # APPENDIX
    doc.add_heading("Appendix: Design Links",1)
    table_with_header(doc,["Design","Platform","Link"],[
        ("Hi-Fi Designs (Phase 4)","Google Stitch","https://stitch.withgoogle.com/projects/12363635124012461076"),
        ("Lo-Fi Wireframes (Phase 3)","Google Stitch","https://stitch.withgoogle.com/projects/14852019868932687382"),
        ("Alternative Design 1","Figma","https://www.figma.com/make/f8gpGhvqfnCadzPZSdnxCw/alterntive-1"),
    ],[2.0,1.5,3.0])

    doc.save(OUT)
    sz = os.path.getsize(OUT)/1024
    print(f"Phase 4 Report generated: {OUT} ({sz:.0f} KB)")

if __name__=="__main__": build()
