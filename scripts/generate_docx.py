"""
PlaySmart Phase 3 Report – DOCX Generator
Generates a downloadable, editable .docx with all photos and hyperlinks.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
STORYBOARD_DIR = os.path.join(BASE_DIR, "storyboard-images")
WIREFRAME_DIR = os.path.join(BASE_DIR, "wireframe-images")
OUTPUT_FILE = os.path.join(BASE_DIR, "PlaySmart-Phase3-Report.docx")

# ─── Helpers ────────────────────────────────────────────────────────────────

def add_hyperlink(paragraph, url, text, color="0563C1", bold=False):
    """Insert a clickable hyperlink into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = parse_xml(f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>')
    run_elem = parse_xml(
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:rPr>'
        f'<w:color w:val="{color}"/>'
        f'<w:u w:val="single"/>'
        f'{"<w:b/>" if bold else ""}'
        f'</w:rPr>'
        f'<w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
    )
    hyperlink.append(run_elem)
    paragraph._element.append(hyperlink)
    return hyperlink


def set_cell_shading(cell, color):
    """Apply background shading to a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    tc_pr.append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a formatted table with header row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2E75B6")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[1 + r_idx].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D9E2F3")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def add_image_safe(doc, path, width=None, caption=None):
    """Insert image if it exists, otherwise show placeholder text."""
    if os.path.exists(path):
        try:
            img = Image.open(path)
            w, h = img.size
            # Default width is 5.5 inches max
            target_w = width or min(5.5, w / 96)
            doc.add_picture(path, width=Inches(target_w))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            p = doc.add_paragraph(f"[Image could not be loaded: {os.path.basename(path)}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f"[Image not found: {os.path.basename(path)}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_wireframe_pair(doc, img1_path, label1, img2_path, label2, width=2.6):
    """Insert two wireframe images side-by-side in a table."""
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove borders
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'</w:tcBorders>'
            )
            tc_pr.append(borders)

    # Images row
    for i, (img_path, label) in enumerate([(img1_path, label1), (img2_path, label2)]):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists(img_path):
            try:
                run = p.add_run()
                run.add_picture(img_path, width=Inches(width))
            except Exception:
                p.add_run(f"[Image: {os.path.basename(img_path)}]")
        else:
            p.add_run(f"[Image not found: {os.path.basename(img_path)}]")

    # Labels row
    for i, label in enumerate([label1, label2]):
        cell = table.rows[1].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def section_heading(doc, text, level=1):
    """Add a section heading."""
    h = doc.add_heading(text, level=level)
    return h


# ─── Main Document Builder ─────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # ===========================================================================
    # TITLE PAGE
    # ===========================================================================
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PlaySmart")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x97)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Phase 3: Low-Fidelity Design and Interaction Planning")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph("")

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("A Human-Centered Football Training and Performance Planning Interface")
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")
    doc.add_paragraph("")

    for line in [
        "Team Members: Usman Wajid, Ahmad Masood, Abdulrehman Naseer",
        "Instructor: Arslan Asif",
        "Date: April 2025",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ===========================================================================
    # TABLE OF CONTENTS
    # ===========================================================================
    section_heading(doc, "Table of Contents", level=1)

    toc_items = [
        "1. Introduction",
        "2. Low-Fidelity Prototypes",
        "    2.1 Storyboards",
        "    2.2 Wireframes – Main Design",
        "    2.3 Wireframes – Alternative Design",
        "3. Design Goals",
        "    3.1 Usability Goals (Quantitative)",
        "    3.2 Usability Goals (Qualitative)",
        "    3.3 User Experience (UX) Goals",
        "4. User Evaluation Activity",
        "    4.1 Overall Feedback",
        "    4.2 Usability Issues Identified",
        "    4.3 Wireframe Comparison",
        "    4.4 Improvements for Next Iteration",
        "5. Conclusion",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ===========================================================================
    # 1. INTRODUCTION
    # ===========================================================================
    section_heading(doc, "1. Introduction", level=1)

    doc.add_paragraph(
        "This phase builds on the research and requirements gathered in Phases 1 and 2 of the PlaySmart project. "
        "The objective of this phase is to translate the requirements, personas, and scenarios into early design "
        "representations that can be evaluated and refined before moving to higher-fidelity designs."
    )
    doc.add_paragraph(
        "PlaySmart is a conceptual mobile-based interface designed to support football players and coaches in planning, "
        "organizing, and reviewing training activities through intuitive and user-friendly interaction. The system focuses "
        "on simplifying training management while reducing cognitive load for users."
    )

    p = doc.add_paragraph("In this phase, we present:")
    bullets = [
        "Storyboards illustrating key user interactions",
        "Wireframes for major system screens (two alternative designs)",
        "Defined usability and UX goals to guide future evaluation",
        "User evaluation feedback from the assigned user group",
    ]
    for b in bullets:
        bp = doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph(
        "The wireframes have been created as low-fidelity grayscale prototypes following standard HCI wireframing rules "
        "(no color, image placeholders, standard UI symbols):"
    )

    # Links
    p = doc.add_paragraph()
    run = p.add_run("Main Design (Lo-Fi Wireframes): ")
    run.bold = True
    add_hyperlink(p, "https://stitch.withgoogle.com/projects/14852019868932687382",
                  "PlaySmart Lo-Fi Wireframes – Stitch")

    p = doc.add_paragraph()
    run = p.add_run("Alternative Design: ")
    run.bold = True
    run2 = p.add_run("Described in Section 2.3 below")

    doc.add_page_break()

    # ===========================================================================
    # 2. LOW-FIDELITY PROTOTYPES
    # ===========================================================================
    section_heading(doc, "2. Low-Fidelity Prototypes", level=1)

    # ── 2.1 Storyboards ──
    section_heading(doc, "2.1 Storyboards (Early Low-Fidelity Prototypes)", level=2)
    doc.add_paragraph(
        "Storyboards visually illustrate how users interact with the PlaySmart system in different situations. "
        "Each storyboard represents a sequence of steps showing how a user performs a task using the system."
    )

    # Storyboard 1
    section_heading(doc, "Storyboard 1: Player Planning a Solo Training Session", level=3)
    doc.add_paragraph(
        "This storyboard illustrates Ahmad's journey from unstructured training to guided, "
        "app-assisted practice using PlaySmart."
    )

    storyboard_frames = [
        ("frame1.png", "Frame 1 — The Problem: Arriving Without Direction",
         "Ahmad, a 21-year-old university football player, arrives at the football ground carrying his ball. "
         "He wants to improve his skills before an upcoming tournament but feels lost — he doesn't know which "
         "drills to practice or how to structure his training session."),
        ("frame2.png", "Frame 2 — Unstructured Training Environment",
         "At the ground, Ahmad sees other players training in a completely random and unstructured manner — "
         "some kicking balls aimlessly, some jogging, others standing around. There's no organized training plan, "
         "no guidance, and no measurable progress being made."),
        ("frame3.png", "Frame 3 — Discovering PlaySmart",
         "Frustrated, Ahmad sits down on a bench beside the field and pulls out his phone. He opens the PlaySmart app, "
         "which greets him with a clean, intuitive splash screen featuring the app logo and a football icon."),
        ("frame4.png", "Frame 4 — Browsing Training Categories & Selecting Drills",
         "The app displays the Trainings screen with three organized categories: Shooting ⚽, Speed ⚡, and Tactics 🧩. "
         "Ahmad taps on \"Speed\" and sees a structured drill list:\n"
         "1. 10x Jump-Jacks\n2. 10x Sprint\n3. 5x High-Knees\n\nEach drill has clear rep counts and instructions."),
        ("frame5.png", "Frame 5 — Performing the Exercises",
         "Ahmad follows the app's guidance and performs each exercise on the field. He does jump-jacks, sprints, "
         "and high-knees as instructed. The app tracks his progress in real time, and a timer keeps him on schedule "
         "throughout the session."),
        ("frame6.png", "Frame 6 — Training Completed Successfully",
         "After completing all exercises, the app displays a \"Training Completed!\" screen with a green checkmark. "
         "Ahmad sees his session summary — total time, exercises completed, and an encouraging message. He feels "
         "accomplished and motivated to continue using PlaySmart for his daily training."),
    ]

    for filename, frame_title, description in storyboard_frames:
        p = doc.add_paragraph()
        run = p.add_run(frame_title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

        img_path = os.path.join(STORYBOARD_DIR, filename)
        add_image_safe(doc, img_path, width=5.0, caption=frame_title)
        doc.add_paragraph(description)
        doc.add_paragraph("")  # spacer

    doc.add_page_break()

    # Storyboard 2
    section_heading(doc, "Storyboard 2: Coach Organizing a Team Training Session", level=3)
    sb2_rows = [
        ("1", "Context: Hamza, a university football coach, needs to organize a training session focusing on passing accuracy and teamwork for the upcoming match. He usually plans manually, which takes considerable time."),
        ("2", "Action: Hamza opens the PlaySmart app on his tablet. The coach dashboard shows options: \"Create Session,\" \"Drill Library,\" \"Player Roster,\" and \"Session History.\""),
        ("3", "Action: He taps \"Create Session\" and selects the date and training focus (Passing, Teamwork, Small-sided Games)."),
        ("4", "Action: The system filters the drill library to show relevant exercises. Hamza browses through visual drill cards showing drill name, duration, equipment needed, and skill focus area."),
        ("5", "Action: Hamza selects exercises for warm-up (10 min), skill drills (25 min), and a practice match (20 min). The app auto-arranges them into a structured session plan."),
        ("6", "Action: During training, players follow the drill sequence displayed on Hamza's tablet. Each drill card shows instructions, player formations, and objectives."),
        ("7", "Action: After the session, Hamza records quick notes about individual player performance using a simple rating system (1–5 stars) and short text notes."),
        ("8", "Outcome: The system stores session data. Hamza can review player progress over time and adjust future training plans accordingly. The session summary shows total training time, drills completed, and player ratings."),
    ]
    add_styled_table(doc, ["Step", "Scene Description"], sb2_rows, col_widths=[0.6, 5.9])

    doc.add_paragraph("")
    doc.add_page_break()

    # Storyboard 3
    section_heading(doc, "Storyboard 3: Team Preparing for a Match", level=3)
    sb3_rows = [
        ("1", "Context: The university football team has an important match next week. Coach Hamza wants to ensure players are fully prepared. He decides to use PlaySmart's Match Preparation feature."),
        ("2", "Action: Hamza opens PlaySmart and navigates to \"Match Prep.\" The system displays a pre-match checklist: Fitness Check, Tactical Drills, Warm-up Plan, and Mental Readiness."),
        ("3", "Action: He reviews the team's recent training data. The app highlights players with high fatigue levels (shown with red/yellow indicators) who may need lighter training."),
        ("4", "Action: Hamza creates a match-day preparation plan with modified intensity for fatigued players and full training for fit players."),
        ("5", "Action: Ahmad (a player) opens the app on his phone and sees his personalized match-day schedule, including warm-up drills and tactical instructions."),
        ("6", "Action: After the preparation session, Hamza uses the \"Readiness Check\" feature where players self-report their physical and mental readiness (scale 1–10)."),
        ("7", "Outcome: The system generates a team readiness overview showing all players' status. Hamza feels confident the team is prepared, and players appreciate the structured approach."),
    ]
    add_styled_table(doc, ["Step", "Scene Description"], sb3_rows, col_widths=[0.6, 5.9])

    doc.add_page_break()

    # ── 2.2 Wireframes – Main Design ──
    section_heading(doc, "2.2 Wireframes – Main Design", level=2)

    doc.add_paragraph(
        "The main design wireframes follow standard low-fidelity wireframing principles:"
    )
    for item in [
        "No Color: Only grayscale (white, black, and gray)",
        "Placeholders: Rectangles with \"X\" for images, simple lines for text",
        "Standard Symbols: Common UI patterns (magnifying glass for search, chevrons for navigation)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run("Interactive Prototype Link: ")
    run.bold = True
    add_hyperlink(p, "https://stitch.withgoogle.com/projects/14852019868932687382",
                  "PlaySmart Lo-Fi Wireframes – Stitch")

    doc.add_paragraph("")

    section_heading(doc, "Wireframe Screens", level=3)

    # Wireframe pairs
    wireframe_pairs = [
        ("Splash / Welcome Screen & Login Screen",
         "splash.png", "Splash / Welcome Screen",
         "login.png", "Login / Registration"),
        ("Player Dashboard & Training Categories",
         "dashboard.png", "Home Dashboard (Player)",
         "trainings.png", "Training Categories"),
        ("Speed Training Detail & Drill Library",
         "speed-detail.png", "Speed Training Detail",
         "drill-library.png", "Drill Library"),
        ("Training Completed & Player Progress",
         "completed.png", "Training Completed",
         "progress.png", "Player Progress"),
    ]

    for pair_title, img1, lbl1, img2, lbl2 in wireframe_pairs:
        p = doc.add_paragraph()
        run = p.add_run(pair_title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

        add_wireframe_pair(
            doc,
            os.path.join(WIREFRAME_DIR, img1), lbl1,
            os.path.join(WIREFRAME_DIR, img2), lbl2,
            width=2.4
        )
        doc.add_paragraph("")

    # Screen descriptions table
    section_heading(doc, "Screen Descriptions", level=3)
    screen_rows = [
        ("Splash / Welcome Screen", "App logo, tagline \"Train Smarter, Play Better,\" and Get Started button"),
        ("Login / Registration", "Simple login form with email/password fields, social login options, and registration link"),
        ("Home Dashboard (Player)", "Central hub showing quick actions: Plan Training, Drill Library, My Progress, Match Prep. Displays upcoming session card and weekly progress bar"),
        ("Home Dashboard (Coach)", "Coach-specific view with: Create Session, Player Roster, Drill Library, Session History. Shows team overview and next scheduled session"),
        ("Drill Library", "Grid/list view of drills with category filters (Fitness, Skills, Tactical). Each drill card shows: name, duration, difficulty, and equipment"),
        ("Drill Detail", "Full drill description with step-by-step instructions, visual diagram/illustration, duration, and \"Add to Session\" button"),
        ("Training Plan / Calendar", "Weekly calendar view showing scheduled sessions. Users can drag-and-drop drills into time slots"),
        ("Session Creation (Coach)", "Multi-step form: Select Date → Choose Focus → Pick Drills → Review Plan → Create Session"),
        ("Player Progress", "Visual dashboard showing: sessions completed, drills finished, weekly/monthly trends using simple bar charts"),
        ("Workload & Fatigue Tracker", "Manual input screen for perceived effort (RPE scale 1–10). Visual fatigue heatmap showing training load over time"),
        ("Match Preparation", "Pre-match checklist with readiness indicators. Training suggestions based on upcoming match date"),
        ("Session Summary", "Post-training overview: drills completed, total time, player ratings (coach view), and improvement notes"),
        ("Settings / Profile", "User profile, notification preferences, and account management"),
    ]
    add_styled_table(doc, ["Screen", "Description"], screen_rows, col_widths=[2.0, 4.5])

    doc.add_paragraph("")

    # Navigation structure
    section_heading(doc, "Navigation Structure (Main Design)", level=3)
    nav_text = (
        "Bottom Navigation Bar:\n"
        "├── Home (Dashboard)\n"
        "├── Drills (Library)\n"
        "├── Plan (Training Calendar)\n"
        "├── Progress (Tracking)\n"
        "└── Profile (Settings)"
    )
    p = doc.add_paragraph()
    run = p.add_run(nav_text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    doc.add_page_break()

    # ── 2.3 Wireframes – Alternative Design ──
    section_heading(doc, "2.3 Wireframes – Alternative Design 1", level=2)

    p = doc.add_paragraph()
    run = p.add_run("The alternative design wireframes are available in Figma: ")
    add_hyperlink(p, "https://www.figma.com/make/f8gpGhvqfnCadzPZSdnxCw/alterntive-1",
                  "Alternative Design 1 – Figma")

    doc.add_paragraph(
        "The alternative design explores different layout and navigation approaches:"
    )

    # Key differences table
    section_heading(doc, "Key Differences from Main Design", level=3)
    diff_rows = [
        ("Navigation", "Bottom tab bar with 5 tabs", "Side drawer/hamburger menu navigation"),
        ("Dashboard Layout", "Card-based grid layout with quick actions", "List-based layout with expandable sections"),
        ("Drill Library", "Grid view with visual drill cards", "List view with filter sidebar"),
        ("Training Plan", "Calendar-based weekly view", "Timeline-based sequential view"),
        ("Progress Display", "Bar charts and trend lines", "Circular progress rings and milestone badges"),
        ("Coach View", "Separate dashboard tab", "Integrated toggle switch between Player/Coach mode"),
    ]
    add_styled_table(doc, ["Aspect", "Main Design", "Alternative Design 1"], diff_rows, col_widths=[1.5, 2.5, 2.5])

    doc.add_paragraph("")

    # Alternative screen descriptions
    section_heading(doc, "Alternative Screen Descriptions", level=3)
    alt_rows = [
        ("Home Screen", "List-based dashboard with expandable sections for Recent Activity, Upcoming Sessions, and Quick Start options"),
        ("Drill Browser", "Left-panel filter sidebar with main content area showing drill list. Filters include: Category, Difficulty, Duration, Equipment"),
        ("Session Timeline", "Vertical timeline showing past and upcoming sessions with status indicators (completed, upcoming, missed)"),
        ("Progress Rings", "Circular progress indicators for weekly goals: Drills Completed, Training Hours, Skill Improvement"),
        ("Quick Drill Start", "Simplified flow: Select Category → Get Recommended Drill → Start Timer → Mark Complete"),
        ("Team Overview (Coach)", "Grid of player cards showing individual readiness, recent activity, and performance rating"),
    ]
    add_styled_table(doc, ["Screen", "Description"], alt_rows, col_widths=[2.0, 4.5])

    doc.add_page_break()

    # ===========================================================================
    # 3. DESIGN GOALS
    # ===========================================================================
    section_heading(doc, "3. Design Goals", level=1)

    # 3.1 Quantitative
    section_heading(doc, "3.1 Usability Goals (Quantitative)", level=2)
    doc.add_paragraph(
        "These are measurable objectives that define how effectively users can interact with the PlaySmart system:"
    )
    quant_rows = [
        ("Task Completion Time", "Time to create a training session", "≤ 3 minutes for new users; ≤ 1 minute for returning users"),
        ("Task Completion Time", "Time to find and start a specific drill", "≤ 30 seconds"),
        ("Error Rate", "Percentage of user actions resulting in errors", "< 5% of user actions"),
        ("Learnability", "Time for a new user to perform basic tasks without assistance", "≤ 5 minutes"),
        ("Efficiency", "Number of taps/interactions to complete common tasks", "≤ 4 taps to start a drill; ≤ 6 to create session"),
        ("Task Success Rate", "Percentage of users completing core tasks without help", "≥ 90% on first attempt"),
        ("System Usability Scale (SUS)", "Standardized usability questionnaire score", "≥ 72 (above average)"),
    ]
    add_styled_table(doc, ["Goal", "Metric", "Target"], quant_rows, col_widths=[1.8, 2.5, 2.2])

    doc.add_paragraph("")

    # 3.2 Qualitative
    section_heading(doc, "3.2 Usability Goals (Qualitative)", level=2)
    doc.add_paragraph(
        "These describe the subjective quality attributes the system should exhibit:"
    )
    qual_items = [
        ("Ease of Use", "The interface should be simple and intuitive, even for first-time users with limited technical expertise. Players should be able to use the app on the field during training without confusion."),
        ("Clarity of Navigation", "Users should always understand where they are in the app and what actions are available. The navigation structure should follow a clear hierarchy that matches users' mental models of training planning."),
        ("Feedback Visibility", "System status (e.g., \"Session Saved,\" \"Drill Added,\" \"Processing…\") should be clearly communicated through visual feedback (toast notifications, animations, color changes)."),
        ("Consistency", "Similar actions and components should behave consistently across all screens. For example, the drill card component should look and behave the same whether viewed in the Library, in a Session Plan, or in the Progress tracker."),
        ("Error Prevention", "The system should guide users to avoid mistakes. For example: disabling the \"Create Session\" button until all required fields are filled, showing validation hints for invalid inputs, and confirming destructive actions."),
        ("User Satisfaction", "After completing tasks, users should feel a sense of accomplishment and clarity. Post-session summaries and progress updates should reinforce positive training habits."),
    ]
    for title, desc in qual_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph("")

    # 3.3 UX Goals
    section_heading(doc, "3.3 User Experience (UX) Goals", level=2)
    doc.add_paragraph(
        "These describe the emotional and experiential qualities the system should provide:"
    )
    ux_rows = [
        ("Enjoyable", "The app should feel pleasant to use. Clean visuals, smooth interactions, and football-themed design elements should make training planning feel like a positive experience rather than a chore."),
        ("Trustworthy", "Users should feel confident that their data (sessions, progress, workload) is accurately recorded and reliably accessible. The system should appear professional and dependable."),
        ("Engaging", "The interaction should feel responsive and alive. Progress tracking, weekly goals, and drill recommendations should encourage users to return and continue training."),
        ("Motivating", "Visual progress indicators, streak tracking, and session completion summaries should motivate players to maintain consistent training habits. Coaches should feel empowered to make better training decisions."),
        ("Supportive", "The app should feel like a helpful training companion rather than an overwhelming data tool. Guidance, suggestions, and contextual help should make users feel supported throughout their experience."),
        ("Focused", "The system should reduce cognitive load by showing only relevant information at each step. Unnecessary clutter should be avoided, and the most important actions should always be prominent."),
    ]
    add_styled_table(doc, ["UX Goal", "Description"], ux_rows, col_widths=[1.3, 5.2])

    doc.add_page_break()

    # ===========================================================================
    # 4. USER EVALUATION ACTIVITY
    # ===========================================================================
    section_heading(doc, "4. User Evaluation Activity", level=1)

    # 4.1
    section_heading(doc, "4.1 Evaluation Process", level=2)
    doc.add_paragraph(
        "The low-fidelity prototypes (storyboards and wireframes) were shared with the assigned user group acting as "
        "representative end-users and client stakeholders. The evaluation was conducted through a structured review session "
        "where users were asked to:"
    )
    eval_steps = [
        "Walk through the storyboards and describe their understanding of each scenario",
        "Review both wireframe designs and identify strengths and weaknesses",
        "Attempt to mentally complete key tasks using the wireframes (cognitive walkthrough)",
        "Provide feedback on usability, clarity, completeness, and overall experience",
    ]
    for i, step in enumerate(eval_steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    # 4.2
    section_heading(doc, "4.2 Overall Feedback", level=2)
    doc.add_paragraph("Users found the concept of PlaySmart very useful, particularly for:")
    for item in [
        "Organizing individual and team training sessions with clear structure",
        "Having a dedicated drill library with visual categorization",
        "Tracking progress and maintaining training consistency",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "The storyboards were easy to follow and effectively communicated the overall system workflow. "
        "Most users were able to clearly understand all three key scenarios."
    )

    # 4.3
    section_heading(doc, "4.3 Usability Issues Identified", level=2)
    issues_rows = [
        ("1", "Users wanted clearer visual distinction between player and coach views", "Medium", "Dashboard"),
        ("2", "The drill difficulty indicators were not immediately obvious", "Low", "Drill Library"),
        ("3", "First-time users needed guidance for creating a training plan – the calendar view felt slightly overwhelming", "Medium", "Training Plan"),
        ("4", "The fatigue tracking input (RPE scale) lacked context – users weren't sure what numbers 1–10 meant without labels", "Medium", "Workload Tracker"),
        ("5", "No confirmation or undo option after adding drills to a session plan", "Low", "Session Creation"),
        ("6", "The match preparation checklist needed more explanation about what each item involved", "Medium", "Match Prep"),
        ("7", "Users suggested adding a quick-start option to begin training immediately without extensive planning", "High", "Home Dashboard"),
    ]
    add_styled_table(doc, ["#", "Issue", "Severity", "Related Screen"], issues_rows, col_widths=[0.4, 3.5, 0.9, 1.5])

    doc.add_paragraph("")

    # 4.4
    section_heading(doc, "4.4 Wireframe Comparison", level=2)
    comp_rows = [
        ("Navigation", "Bottom tab bar felt familiar and easy to use", "Hamburger menu was less discoverable; users had to learn where things were"),
        ("Visual Clarity", "Card-based layout was visually appealing and easy to scan", "List-based layout was more compact but felt cluttered"),
        ("Drill Library", "Grid view with visual cards made browsing enjoyable and quick", "Filter sidebar was powerful but added complexity for casual users"),
        ("Training Planning", "Calendar view was intuitive for weekly planning", "Timeline view was novel but less familiar; some confusion"),
        ("Progress Tracking", "Bar charts were clear and easy to interpret", "Circular progress rings were more engaging but harder to compare across weeks"),
        ("Overall Preference", "Preferred by 65% of evaluators – felt more intuitive", "Preferred by 35% – felt more information-dense and flexible"),
    ]
    add_styled_table(doc, ["Aspect", "Main Design Feedback", "Alternative Design 1 Feedback"], comp_rows, col_widths=[1.4, 2.5, 2.6])

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Key Insight: ")
    run.bold = True
    p.add_run(
        "The evaluation indicates a need to combine strengths from both designs — maintaining the intuitive navigation "
        "and visual clarity of the Main Design while incorporating the powerful filtering and progress visualization "
        "elements from the Alternative Design."
    )

    # 4.5
    section_heading(doc, "4.5 Improvements for Next Iteration", level=2)
    doc.add_paragraph(
        "Based on the feedback received, the following refinements will be made to improve usability and user experience:"
    )
    improvements = [
        ("Quick-Start Training Feature", "A prominent \"Start Training Now\" button will be added to the home dashboard, allowing users to begin a session with recommended drills in just 2 taps, bypassing full session planning when desired."),
        ("Role-Based Onboarding", "A short onboarding flow will differentiate between player and coach modes, customizing the dashboard and features accordingly. Users can switch roles at any time."),
        ("RPE Scale Labels", "The fatigue tracking input will include descriptive labels alongside numbers (e.g., 1 = \"No Effort,\" 5 = \"Moderate,\" 10 = \"Maximum Effort\") with color-coded visual cues."),
        ("Drill Difficulty Badges", "Clearer difficulty indicators will be added to drill cards using universally understood symbols (e.g., 1–3 stars, or color-coded labels: Green = Beginner, Yellow = Intermediate, Red = Advanced)."),
        ("Contextual Guidance", "Placeholder text, tooltips, and brief instructional hints will be added throughout the interface to support first-time users."),
        ("Confirmation Dialogs", "Undo and confirmation actions will be incorporated for critical operations like removing drills from a session or deleting a training plan."),
        ("Match Prep Enhancement", "The pre-match checklist will include brief descriptions and expandable details for each preparation step to guide coaches through the process."),
        ("Balanced Interface Design", "The next iteration will adopt the Main Design's bottom navigation and card-based layout while integrating the Alternative Design's filter sidebar (as a collapsible overlay) and circular progress elements for weekly goals."),
    ]
    for i, (title, desc) in enumerate(improvements, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {title}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ===========================================================================
    # 5. CONCLUSION
    # ===========================================================================
    section_heading(doc, "5. Conclusion", level=1)
    doc.add_paragraph(
        "This phase successfully translated the initial concept and requirements of PlaySmart into low-fidelity design "
        "representations, including storyboards and wireframes. These prototypes helped visualize key user interactions "
        "and provided a clear understanding of how users will engage with the system."
    )

    p = doc.add_paragraph()
    run = p.add_run("Key Accomplishments:")
    run.bold = True
    for item in [
        "Three comprehensive storyboards illustrating core user scenarios (individual player, coach, and team match preparation)",
        "Two alternative wireframe designs created in Figma, exploring different navigation and layout approaches",
        "Well-defined usability goals (quantitative and qualitative) and UX goals to guide system evaluation",
        "Structured user evaluation with actionable feedback",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Summary of Design Decisions:")
    run.bold = True
    for item in [
        "The Main Design (bottom tab navigation, card-based layout) was preferred by the majority of evaluators and will serve as the primary design direction",
        "Elements from the Alternative Design (advanced filtering, circular progress indicators) will be selectively incorporated",
        "Eight specific improvements have been identified based on user feedback",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "The defined usability and UX goals establish a strong foundation for evaluating the effectiveness of the design "
        "in subsequent phases. The user evaluation provided valuable insights into both the strengths and areas for "
        "improvement in the current design, ensuring that the next iteration addresses real user needs."
    )
    doc.add_paragraph(
        "Overall, this phase plays a crucial role in identifying potential usability issues early, allowing for informed "
        "refinements before moving to high-fidelity prototypes and implementation."
    )

    # ===========================================================================
    # APPENDIX
    # ===========================================================================
    doc.add_page_break()
    section_heading(doc, "Appendix: Design Links", level=1)

    link_table = doc.add_table(rows=3, cols=3)
    link_table.style = "Table Grid"
    link_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    headers = ["Design", "Platform", "Link"]
    for i, h in enumerate(headers):
        cell = link_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2E75B6")

    # Row 1
    link_table.rows[1].cells[0].text = "Main Design (Lo-Fi Wireframes)"
    link_table.rows[1].cells[1].text = "Google Stitch"
    cell = link_table.rows[1].cells[2]
    cell.text = ""
    p = cell.paragraphs[0]
    add_hyperlink(p, "https://stitch.withgoogle.com/projects/14852019868932687382",
                  "https://stitch.withgoogle.com/projects/14852019868932687382")

    # Row 2
    link_table.rows[2].cells[0].text = "High-Fidelity Prototype (Reference)"
    link_table.rows[2].cells[1].text = "Google Stitch"
    cell = link_table.rows[2].cells[2]
    cell.text = ""
    p = cell.paragraphs[0]
    add_hyperlink(p, "https://stitch.withgoogle.com/projects/12363635124012461076",
                  "https://stitch.withgoogle.com/projects/12363635124012461076")

    for row in link_table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    # Save
    doc.save(OUTPUT_FILE)
    print(f"\nDOCX generated successfully: {OUTPUT_FILE}")
    print(f"File size: {os.path.getsize(OUTPUT_FILE) / 1024:.1f} KB")

if __name__ == "__main__":
    build_document()
