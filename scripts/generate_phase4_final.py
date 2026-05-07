"""PlaySmart Phase 4 Report — Final PDF Generator
Matches CodeMap Phase 4 format exactly:
  Title > Introduction > Typography > Color > Grouping > Navigation > UI Elements > Realistic Content
Each section has Status, description, and relevant screenshot(s).
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

BASE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(BASE, "hifi-screenshots", "individual")
OUT = os.path.join(BASE, "PlaySmart-Phase4-Report.docx")

# ─── helpers ───────────────────────────────────────────────────────────

def add_hyperlink(p, url, text):
    part = p.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl = parse_xml(f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>')
    r = parse_xml(f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r>')
    hl.append(r); p._element.append(hl)

def shade(cell, color):
    tc = cell._element.get_or_add_tcPr()
    tc.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'))

def table_with_header(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; shade(c,"2E75B6")
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c = t.rows[1+ri].cells[ci]; c.text = v
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
            if ri%2==1: shade(c,"D9E2F3")
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    return t

def add_img(doc, name, w=4.5, caption=None):
    path = os.path.join(SHOTS, f"{name}.png")
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f"[Image: {name}.png not found]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].italic=True; c.runs[0].font.size=Pt(9); c.runs[0].font.color.rgb=RGBColor(0x66,0x66,0x66)

def status_badge(doc, status="Met"):
    p = doc.add_paragraph()
    r = p.add_run(f"Status: {status}")
    r.bold = True
    if status == "Met":
        r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    elif status == "Partially Met":
        r.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
    else:
        r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

def section_heading(doc, num, title):
    doc.add_heading(f"{num} {title}", 1)

# ─── Build ─────────────────────────────────────────────────────────────

def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════════════
    for _ in range(6): doc.add_paragraph("")
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PlaySmart"); r.bold = True; r.font.size = Pt(36); r.font.color.rgb = RGBColor(0xCA,0xFD,0x00)
    t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Phase 4: High-Fidelity Design and Evaluation"); r2.font.size = Pt(18); r2.font.color.rgb = RGBColor(0x44,0x44,0x44)
    doc.add_paragraph("")
    for line in [
        "Human Computer Interaction",
        "",
        "Team Members: Usman Wajid, Ahmad Masood, Abdulrehman Naseer",
        "Instructor: Arslan Asif",
        "Date: April 2026",
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x55,0x55,0x55)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ═══════════════════════════════════════════════════════════════════
    section_heading(doc, "1.", "Introduction")
    doc.add_paragraph(
        "PlaySmart is a football training management platform designed to help individual players "
        "and coaches organize, track, and optimize their training sessions. It enables users to "
        "browse drill libraries, plan structured training sessions, track performance progress, "
        "and prepare for matches \u2014 all through a cohesive mobile interface."
    )
    doc.add_paragraph(
        "This document presents the Phase 4 deliverables for the Human Computer Interaction course. "
        "It focuses on a high-fidelity design assessment of the PlaySmart prototype, evaluating "
        "both its visual design and interaction quality. In addition, a formal heuristic evaluation "
        "has been conducted using Nielsen\u2019s 10 Usability Heuristics to systematically identify "
        "usability strengths and areas for improvement. Together, these analyses aim to provide a "
        "comprehensive understanding of the platform\u2019s user experience and design effectiveness."
    )
    doc.add_paragraph("")
    p = doc.add_paragraph()
    r = p.add_run("Interactive Prototype: "); r.bold = True
    add_hyperlink(p, "https://stitch.withgoogle.com/projects/12363635124012461076", "https://stitch.withgoogle.com/projects/12363635124012461076")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 2. HIGH-FIDELITY DESIGN
    # ═══════════════════════════════════════════════════════════════════
    section_heading(doc, "2.", "High-Fidelity Design")
    doc.add_paragraph(
        "This section evaluates the PlaySmart prototype against each requirement listed in the "
        "Phase 4 assignment brief. Each criterion is assessed as Met, Partially Met, or Not Met, "
        "with observations grounded in the actual screens produced."
    )
    p = doc.add_paragraph()
    r = p.add_run("Stitch link: "); r.bold = True
    add_hyperlink(p, "https://stitch.withgoogle.com/projects/12363635124012461076", "https://stitch.withgoogle.com/projects/12363635124012461076")
    doc.add_paragraph("")

    # ─── 2.1 Typography ───────────────────────────────────────────────
    doc.add_heading("2.1 Typography", 2)
    status_badge(doc, "Met")
    doc.add_paragraph(
        "The design consistently uses a modern, clean typeface combination across all 14 screens, "
        "creating a cohesive and professional visual identity. Space Grotesk is used for display "
        "headings and large performance statistics, giving the interface a bold, athletic feel. "
        "Inter is used for all body text, labels, and secondary content, ensuring excellent "
        "readability at smaller sizes."
    )
    doc.add_paragraph(
        "A clear typographic hierarchy is maintained throughout: large H1 headings for screen "
        "titles (e.g., \u201CPlayer Dashboard\u201D, \u201CSpeed Training\u201D), well-defined H2 section "
        "headers for organizing content within screens, and smaller body text for descriptions, "
        "drill instructions, and metadata. All-caps labels are used consistently for metadata "
        "tags (REPS, SETS, BPM, DURATION), creating a clear visual distinction between content "
        "types."
    )
    doc.add_paragraph(
        "An especially strong design decision is the use of contrasting font weights \u2014 bold "
        "headings against regular body text \u2014 which guides the eye through the information "
        "hierarchy naturally. The volt-green (#CAFD00) accent color on key text elements "
        "(like CTA labels) further reinforces the visual hierarchy without relying solely on "
        "font size."
    )
    add_img(doc, "welcome", 4.0, "Figure 1: Welcome Screen \u2013 Typography hierarchy with Space Grotesk headings")
    doc.add_paragraph("")
    add_img(doc, "speed_detail", 4.0, "Figure 2: Speed Training Detail \u2013 All-caps labels for REPS/SETS metadata")
    doc.add_page_break()

    # ─── 2.2 Color ────────────────────────────────────────────────────
    doc.add_heading("2.2 Color", 2)
    status_badge(doc, "Met")
    doc.add_paragraph(
        "A coherent and purposeful color palette is applied consistently across all screens, "
        "contributing to a clear and visually stable interface. The design uses a dark mode "
        "foundation (#0E0E0E charcoal background) paired with high-contrast light text (#FFFFFF), "
        "ensuring strong readability and reducing visual strain during extended use \u2014 particularly "
        "important for athletes checking their phones in bright outdoor conditions."
    )
    doc.add_paragraph(
        "A distinct volt-green (#CAFD00) is used as the primary interactive color and is "
        "consistently applied to key call-to-action elements such as \u201CGet Started,\u201D "
        "\u201CStart Training,\u201D and \u201CEnter Training Facility,\u201D as well as active navigation "
        "states and progress indicators. This consistency helps users quickly identify interactive "
        "components and reinforces predictable behavior across the interface."
    )
    doc.add_paragraph(
        "Semantic color usage is handled effectively to communicate system states. Green indicators "
        "are used for completion status and successful actions. Orange/amber communicates "
        "in-progress or warning states (e.g., medium fatigue levels in Match Preparation). "
        "Red (#FF7351) is reserved for error states, destructive actions, and high-fatigue alerts. "
        "This disciplined use of color strengthens feedback clarity and aligns well with "
        "established usability conventions."
    )
    add_img(doc, "player_dashboard", 4.0, "Figure 3: Player Dashboard \u2013 Volt-green CTAs on dark background")
    doc.add_paragraph("")
    add_img(doc, "ui_states", 4.0, "Figure 4: UI States \u2013 Semantic color: green (success), red (error), orange (warning)")
    doc.add_page_break()

    # ─── 2.3 Grouping / Screen Layout ─────────────────────────────────
    doc.add_heading("2.3 Grouping", 2)
    status_badge(doc, "Met")
    doc.add_paragraph("All five principles of screen design and layout are actively applied across the prototype:")
    doc.add_paragraph("")

    principles = [
        ("Grouping", "Related elements are grouped into distinct cards and sections. The Player Dashboard organizes quick-action cards in a 2\u00d72 grid. The Settings screen groups options into distinct categories (Notifications, Preferences, Account). The Match Preparation screen groups checklist items and team readiness as separate visual blocks."),
        ("Ordering", "Forms and content follow natural reading order from top to bottom. The Training Detail screen positions drill name above description above exercise list above action button. The Drill Library shows category filters above the search bar above drill results."),
        ("Alignment", "Left-alignment is consistently used for content-heavy screens (Dashboard, Drill Library, Progress). Center-alignment is used for the Welcome screen\u2019s hero text and the Training Completed screen\u2019s celebratory summary to create visual focus."),
        ("Whitespace", "Generous internal padding is applied within all cards and sections, preventing visual clutter. The dark background with tonal layering (#0E0E0E \u2192 #1A1A1A \u2192 #262626) creates visual breathing room between sections without explicit divider lines."),
        ("Decoration", "Decoration is minimal and purposeful. The volt-green accent (#CAFD00) is used only for interactive elements and active states. No gratuitous gradients or drop shadows are applied. Status indicator colors (green/orange/red) serve as functional, not decorative, signals."),
    ]
    for title, desc in principles:
        p = doc.add_paragraph()
        r = p.add_run(f"{title}: "); r.bold = True
        p.add_run(desc)

    add_img(doc, "training_categories", 4.0, "Figure 5: Training Categories \u2013 Grouped category cards with consistent spacing")
    doc.add_paragraph("")
    add_img(doc, "profile_settings", 4.0, "Figure 6: Settings \u2013 Grouped sections for Notifications, Preferences, Account")
    doc.add_page_break()

    # ─── 2.4 Navigation Paths ─────────────────────────────────────────
    doc.add_heading("2.4 Navigation Paths", 2)
    status_badge(doc, "Met")
    doc.add_paragraph(
        "All navigational paths within the system are covered by the prototype. The complete user "
        "journey from first visit through to active usage is represented:"
    )

    nav_paths = [
        "Welcome Screen \u2192 Login / Sign Up",
        "Login \u2192 [correct credentials] \u2192 Player Home Dashboard",
        "Dashboard \u2192 Training Categories \u2192 Speed/Shooting/Tactics \u2192 Drill Detail \u2192 Start Training \u2192 Training In Progress \u2192 Training Completed",
        "Dashboard \u2192 Drill Library \u2192 Filter/Search \u2192 Drill Card",
        "Dashboard \u2192 My Progress (via bottom nav) \u2192 Weekly/Monthly stats",
        "Dashboard \u2192 Match Preparation \u2192 Checklist + Team Readiness",
        "Coach Dashboard \u2192 Create Session \u2192 Training Calendar",
        "Profile & Settings \u2192 [save changes] \u2192 Success Toast \u2192 Dashboard",
    ]
    for nav in nav_paths:
        doc.add_paragraph(nav, style="List Bullet")

    doc.add_paragraph(
        "All inner screens include a persistent bottom navigation bar with 5 tabs (Home, Drills, "
        "Plan, Progress, Profile), ensuring no dead ends exist within the prototype flow. "
        "The active tab is highlighted with the primary volt-green color."
    )
    add_img(doc, "drill_library", 4.0, "Figure 7: Drill Library \u2013 Bottom navigation bar visible on all screens")
    doc.add_paragraph("")
    add_img(doc, "coach_dashboard", 4.0, "Figure 8: Coach Dashboard \u2013 Role-specific navigation with all paths accessible")
    doc.add_page_break()

    # ─── 2.5 UI Elements ─────────────────────────────────────────────
    doc.add_heading("2.5 UI Elements", 2)
    status_badge(doc, "Met")
    doc.add_paragraph("All required UI elements are present and integrated into the prototype:")
    doc.add_paragraph("")

    ui_elements = [
        ("Dialog Boxes", "The \u2018Delete Training Session?\u2019 confirmation dialog is a fully designed modal overlay with a title, descriptive warning text, and two action buttons (\u201CCancel\u201D and \u201CDelete\u201D). The dialog uses the error color (#FF7351) for the destructive action to clearly indicate risk."),
        ("Progress Indicators", "The Training In Progress screen displays a real-time circular timer with current exercise name and progress. The Player Progress screen shows a 75% circular progress ring and a daily activity bar chart. The Profile/Settings screen includes a \u201CSyncing training data... 65%\u201D linear progress bar."),
        ("Message/Toast Notifications", "Three distinct toast notifications are designed: (1) green success toast for \u201CSession saved successfully!\u201D, (2) red error toast for \u201CFailed to sync data. Please check your connection,\u201D and (3) informational feedback for completed actions. These serve as the system\u2019s primary feedback mechanism for user-triggered actions."),
    ]
    for title, desc in ui_elements:
        p = doc.add_paragraph()
        r = p.add_run(f"{title}: "); r.bold = True
        p.add_run(desc)

    add_img(doc, "training_inprogress", 4.0, "Figure 9: Training In Progress \u2013 Circular timer and exercise progress indicator")
    doc.add_paragraph("")
    add_img(doc, "ui_states", 4.0, "Figure 10: UI States \u2013 Dialog box, success/error toasts, and progress bar")
    doc.add_page_break()

    # ─── 2.6 Realistic Content ────────────────────────────────────────
    doc.add_heading("2.6 Realistic Content", 2)
    status_badge(doc, "Met")
    doc.add_paragraph(
        "Placeholder text such as \u2018Lorem Ipsum\u2019 is entirely absent from the prototype. "
        "All content is realistic, meaningful, and domain-appropriate for a football training platform:"
    )

    content_examples = [
        "Player names: Ahmad Masood, Coach Arslan",
        "Training categories: Shooting, Speed, Tactics",
        "Drill names: Sprint Intervals, Ball Control Mastery, Passing Triangle, Agility Ladder",
        "Exercise details: 10\u00d7 Jump-Jacks, 10\u00d7 Sprint, 5\u00d7 High-Knees",
        "Progress stats: 12 Sessions, 8.5 Hours, 47 Drills, 75% weekly goal",
        "Match prep items: Fitness Check, Tactical Drills, Warm-up Plan, Mental Readiness",
        "Fatigue indicators: Low (green), Medium (orange), High (red)",
        "Calendar events: Speed Training (Mon 9AM), Team Tactics (Wed 3PM)",
    ]
    for item in content_examples:
        doc.add_paragraph(item, style="List Bullet")

    add_img(doc, "match_preparation", 4.0, "Figure 11: Match Preparation \u2013 Realistic checklist and team readiness data")
    doc.add_paragraph("")
    add_img(doc, "progress_tracking", 4.0, "Figure 12: Progress Tracking \u2013 Realistic stats (12 sessions, 8.5h, 47 drills)")
    doc.add_paragraph("")
    add_img(doc, "calendar", 4.0, "Figure 13: Training Calendar \u2013 Realistic weekly schedule with actual session names")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 3. ALL SCREENS OVERVIEW
    # ═══════════════════════════════════════════════════════════════════
    section_heading(doc, "3.", "Complete Screen Catalog")
    doc.add_paragraph("Below is the complete set of 14 high-fidelity screens designed for PlaySmart:")
    doc.add_paragraph("")

    all_screens = [
        ("welcome", "1. Welcome / Splash Screen"),
        ("login", "2. Login / Sign Up"),
        ("player_dashboard", "3. Player Home Dashboard"),
        ("training_categories", "4. Training Categories"),
        ("speed_detail", "5. Speed Training Detail"),
        ("training_inprogress", "6. Training In Progress"),
        ("training_completed", "7. Training Completed"),
        ("drill_library", "8. Drill Library"),
        ("progress_tracking", "9. Player Progress Tracking"),
        ("coach_dashboard", "10. Coach Dashboard"),
        ("match_preparation", "11. Match Preparation"),
        ("calendar", "12. Training Plan Calendar"),
        ("profile_settings", "13. Profile & Settings"),
        ("ui_states", "14. UI Interaction States"),
    ]

    for fname, title in all_screens:
        add_img(doc, fname, 3.5, title)
        doc.add_paragraph("")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # APPENDIX
    # ═══════════════════════════════════════════════════════════════════
    section_heading(doc, "", "Appendix: Design Links")
    table_with_header(doc, ["Design", "Platform", "Link"], [
        ("Hi-Fi Prototype (Phase 4)", "Google Stitch", "https://stitch.withgoogle.com/projects/12363635124012461076"),
        ("Lo-Fi Wireframes (Phase 3)", "Google Stitch", "https://stitch.withgoogle.com/projects/14852019868932687382"),
    ], [2.0, 1.5, 3.0])

    doc.save(OUT)
    sz = os.path.getsize(OUT) / 1024
    print(f"\nPhase 4 Report generated: {OUT} ({sz:.0f} KB)")
    print(f"Total pages: ~15-18 pages with screenshots")

if __name__ == "__main__":
    build()
